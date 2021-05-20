from json.decoder import JSONDecodeError
from math import inf
from datetime import date
from os import read, replace
from random import seed
from sqlalchemy import engine

from sqlalchemy.orm.exc import UnmappedInstanceError
from sqlalchemy.sql.expression import delete, insert, update
from sqlalchemy.sql.functions import ReturnTypeFromArgs
from training.client.choice_funcs import get_digit_choice, get_yes_no_choice, get_day, get_month, get_year
from training.sock_utils import send_message, decode_message_chunks, get_data_from_connection
from docx import Document
import json
import logging
import socket
import click

def dump_to_json(object, out_file_name):
    try:
        with open(out_file_name, 'w') as f:
            json.dump(object, f, indent=4, sort_keys=True)
    except JSONDecodeError:
        msg = f"JSONDecodeError: There was a problem dumping the JSON object {object} to {out_file_name}"
        print(msg)
        logging.error(msg)

def dump_to_docx(object, out_file_name):
    document = Document()
    document.add_heading("Query Results", 0)
    if isinstance(object, dict):
        # only have one row worth of data
        table = document.add_table(rows=1, cols=len(object))
        table.autofit = True
        # iterate over keys for header cells and vals for data cells
        hdr_cells = table.rows[0].cells
        row_cells = table.add_row().cells
        for i, key, value in enumerate(object.items()):
            hdr_cells[i].text = key
            row_cells[i].text = str(value)
    else:
        # iterate over object for many rows of data
        table = document.add_table(rows=1, cols=len(object[0]))
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(object[0]):
            hdr_cells[i].text = key
        for dic in object:
            row_cells = table.add_row().cells
            for i, value in enumerate(dic.values()):
                row_cells[i].text = str(value)
    document.add_page_break()
    document.save(out_file_name)


class Client:

    def __init__(self, server_port, listen_port):
        self.server_port = server_port
        self.port = listen_port
        self.sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        self.sock.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
        self.sock.bind(("localhost", self.port))
        self.sock.listen()
        self.sock.settimeout(1)
        logging.basicConfig(filename="client.log", level=logging.DEBUG, format="%(asctime)s - %(levelname)s: %(message)s")

    def get_server_response(self):
        server_response = None
        while server_response is None:
            message_chunks = get_data_from_connection(self.sock)

            if not message_chunks:
                continue

            try:
                message_dict = decode_message_chunks(message_chunks)
                server_response = message_dict
                break
            except JSONDecodeError:
                continue
        return server_response

    def check_server_status(self, server_response):
        """Returns False in case of an error. Returns the status otherwise."""
        try:
            status = server_response["status"]
        except:
            error_msg = "Server response did not include entry \"status\" to let the client know how to proceed."
            logging.error(error_msg)
            print(error_msg)
            return False

        if status == "error":
            error_msg = server_response["text"]
            logging.error(error_msg)
            print(error_msg)
            return False

        return status

    # Add a vehicle to the DB, update quantity if model already exists
    def add_vehicle(self, model=None, prompt_engin_assign=True):
        if model is None:
            model = input("Enter the model of the new vehicle:")
        quantity = get_digit_choice("Enter the quantity of new vehicles:",
                                    "Please enter a non-negative quantity.",
                                    0, inf)
        price = get_digit_choice("Enter the price of the new vehicle:",
                                 "Invalid price. Please enter a price > 0",
                                 1, inf)
        year = get_year("Enter the year the new vehicle was manufactured:")
        month = get_month("Enter the month the new vehicle was manufactured:")
        day = get_day("Enter the date the new vehicle was manufactured:")
        manufacture_date = date(year, month, day)

        # Send insert message to the server
        logging.info(f"Asking server to add new vehicle {model} manufactured on {manufacture_date} to the database.")
        insert_msg = {
            "port": self.port,
            "action": "add",
            "data_type": "vehicle",
            "model": model,
            "quantity": quantity,
            "price": price,
            "manufacture_year": year,
            "manufacture_month": month,
            "manufacture_date": day
        }
        send_message("localhost", self.server_port, insert_msg)

        # Wait for server response
        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return

        if status == "updated":
            update_msg = server_response["text"]
            logging.info(update_msg)
            print(update_msg)
            return

        # Else: insert was successful
        new_server_port = server_response["port"]
        # Added a new vehicle, prompt user to assign engineers to the vehicle
        success_msg = f"Successfully added new vehicle!\nModel: {model}\nQuantity: {quantity}\nPrice: {price}\nManufacture Date: {manufacture_date}"
        print(success_msg)
        logging.info(success_msg)

        if prompt_engin_assign:
            assign_engin = get_yes_no_choice("Assign engineers to this vehicle? (Y/N):")
            if assign_engin == 'y':
                engineer_names = input("\nEnter names of engineers to assign.\nSeparate names by a comma (e.g. John Smith, Jane Doe):")
                engineer_names = engineer_names.split(',')
                logging.info(f"Asking server to assign engineers {engineer_names} to the {model} manufactured on {manufacture_date}.")
                assign_msg = {
                    "response": "y",
                    "engineers": engineer_names
                }
                send_message("localhost", new_server_port, assign_msg)

                server_response = self.get_server_response()
                status = self.check_server_status(server_response)
                if not status:
                    return

                assigned = server_response["assigned"]
                unassigned = server_response["unassigned"]

                print("\nSuccessfully assigned engineers to the new vehicle!")
                logging.info(f"\nSuccessfully assigned engineers {assigned} to vehicle model {model} manufactured on {manufacture_date}")
                logging.info(f"\nEngineers {unassigned} did not exist in the database, and could not be assigned to the new vehicle.")
            else:
                assign_msg = {
                    "response": "n"
                }
                send_message("localhost", new_server_port, assign_msg)
                logging.info(f"User skipped assigning any engineers to new vehicle {model} manufactured on {manufacture_date}")

    # Add an engineer to the DB
    def add_engineer(self, engin_name=None, prompt_vehicle_assign=True):
        if engin_name is None:
            engin_name = input("Enter the new Engineer's name:")
        birth_year = get_year(f"Enter engineer {engin_name}'s birth year:")
        birth_month = get_month(f"Enter engineer {engin_name}'s birth month:")
        birth_day = get_day(f"Enter engineer {engin_name}'s birth date:")
        date_of_birth = date(birth_year, birth_month, birth_day)
        

        logging.info(f"Asking server to add new engineers {engin_name} born on {date_of_birth} to the database.")
        insert_msg = {
            "data_type": "engineer",
            "action": "add",
            "port": self.port,
            "name": engin_name,
            "birth_year": birth_year,
            "birth_month": birth_month,
            "birth_date": birth_day
        }
        send_message("localhost", self.server_port, insert_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return None

        try:
            new_server_port = server_response["port"]
        except:
            no_port = "Server response did not include entry \"port\" to let the client know where to send yes/no response."
            logging.error(no_port)
            print(no_port)
            return None

        success_msg = f"Successfully added new engineer {engin_name} to the database!"
        logging.info(success_msg)

        
        if prompt_vehicle_assign:
            assign_vehicle = get_yes_no_choice("Assign the new engineer to any vehicles? (Y/N):")
            if assign_vehicle == 'y':
                vehicle_models = input("Enter models of vehicles to assign.\nIt is assumed that manufacture date is irrelevant.\nSeparate model names by a comma (e.g. Fusion, Bronco):")
                vehicle_models = vehicle_models.split(',')
                logging.info(f"Asking the server to assign engineer {engin_name} to vehicles {vehicle_models}")
                assign_msg = {
                    "response": "y",
                    "vehicles": vehicle_models
                }
                send_message("localhost", new_server_port, assign_msg)

                server_response = self.get_server_response()
                status = self.check_server_status(server_response)

                if not status:
                    return
                
                assigned = server_response["assigned"]
                unassigned = server_response["unassigned"]

                print(f"Successfully assigned {engin_name} to vehicles!")
                logging.info(f"Successfully assigned engineer {engin_name} to vehicles {assigned}.")
                logging.info(f"Vehicle models {unassigned} did not exist in the database, and could not be assigned to the new engineer {engin_name}.")
            else:
                assign_msg = {
                    "response" : "n"
                }
                logging.info(f"User skipped assigning engineer {engin_name} to any vehicles.")
                send_message("localhost", new_server_port, assign_msg)

    # Add a laptop to the DB and loan to an engineer if desired
    def add_laptop(self):
        engin_name = input("Enter the name of the engineer this laptop will be loaned to:")
        engin_name = engin_name.strip()
        model = input("Enter the model of the new laptop:")
        year_loaned = get_year(f"Enter the year the laptop was loaned to {engin_name}:")
        month_loaned = get_month(f"Enter the month the laptop was loaned to {engin_name}:")
        day_loaned = get_day(f"Enter the date the laptop was loaned to {engin_name}")
        logging.info(f"Asking server to add a new laptop to the database. Attempting to loan it to engineer {engin_name}")

        insert_msg = {
            "data_type": "laptop",
            "action": "add",
            "port": self.port,
            "model": model,
            "loan_year": year_loaned,
            "loan_month": month_loaned,
            "loan_date": day_loaned,
            "engineer": engin_name
        }
        send_message("localhost", self.server_port, insert_msg)

        success = False
        while not success:
            server_response = self.get_server_response()

            status = self.check_server_status(server_response)

            if not status:
                return

            elif status == "success":
                success = True
                loaned_to = server_response["engineer"]
                success_msg = f"Successfully added laptop {model} to the database and loaned it to {loaned_to}."
                logging.info(success_msg)
                print(success_msg)
                return

            try:
                new_server_port = server_response["port"]
            except:
                error_msg = "Server response has no entry for \"port\" to let the client know where to send the response to."
                print(error_msg)
                logging.error(error_msg)
                return
            
            if status == "no_engineer":
                logging.info(f"Engineer {engin_name} did not exist in the database.")
                print(f"Engineer {engin_name} did not exist in the database.")
                abort_laptop = get_yes_no_choice(f"Would you still like to add this laptop without loaning it to an engineer? (Y/N):")
                response_msg = {
                    "response": abort_laptop
                }
                send_message("localhost", new_server_port, response_msg)
                if abort_laptop == 'n':
                    logging.info("Client chose to abort adding the laptop without loaning it to an existing engineer.")
                    return
            
            elif status == "previous_laptop":
                warning_msg = f"Engineer {engin_name} already has a laptop loaned to them\n" + \
                          f"\n!!! WARNING !!! - Adding a new laptop would replace the laptop already loaned to {engin_name}" + \
                          f"{engin_name}'s previous laptop would not be deleted from the database, but would be loaned by no one."
                print(warning_msg)
                logging.warning(warning_msg)
                replace_laptop = get_yes_no_choice(f"Replace {engin_name}'s laptop with the new one? (Y/N)")
                response_msg = {
                    "response": replace_laptop
                }
                send_message("localhost", new_server_port, response_msg)
                if replace_laptop == 'n':
                    logging.info(f"Client chose to abort adding the laptop as to not replace {engin_name}'s existing laptop")
                    return
            

    # Add contact details for an engineer
    def add_contact(self):
        engin_name = input("Which engineer's contact info are you providing? Enter their name:")
        engin_name = engin_name.strip()
        phone_number = input(f"Enter {engin_name}'s phone number (XXX)-XXX-XXXX:")
        address = input(f"Enter {engin_name}'s address:")
        logging.info(f"Asking server to add contact details for engineer {engin_name}.")
        insert_msg = {
            "data_type": "contact_details",
            "action": "add",
            "port": self.port,
            "phone_number": phone_number,
            "address": address,
            "engineer": engin_name
        }
        send_message("localhost", self.server_port, insert_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return
        
        success_msg = f"Successfully added new contact information for {engin_name}!"
        print(success_msg)
        logging.info(success_msg)
        return
        
    # Delete an engineer
    def delete_engineer(self):
        engin_name = input("Enter the name of the engineer to be deleted:")
        engin_name = engin_name.strip()
        proceed = get_yes_no_choice(f"!!! WARNING !!! - Deleting engineer {engin_name} will also delete any of their contact details. Proceed? (Y/N):")
        if proceed == 'n':
            print(f"Aborted deleting engineer {engin_name} from the database.")
            logging.info(f"User chose to abort deleting engineer {engin_name} as to not delete their contact details.")
            return
        logging.info(f"Asking server to delete engineer {engin_name} from the database.")
        delete_msg = {
            "data_type": "engineer",
            "action": "delete",
            "port": self.port,
            "name": engin_name
        }
        send_message("localhost", self.server_port, delete_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return
        
        success_msg = f"Successfully deleted engineer {engin_name} and their contact details from the database."
        print(success_msg)
        logging.info(success_msg)

    # Delete a vehicle
    def delete_vehicle(self):
        model = input("Enter the model of the vehicle to delete:")
        model = model.strip()
        logging.info(f"Attempting to delete all {model} vehicle records from the database.")
        proceed = get_yes_no_choice(f"!!! WARNING !!! - ALL {model} vehicle records will be deleted, regardless of manufacture date.\nProceed? (Y/N):")
        if proceed == 'n':
            logging.info(f"User chose to abort deleting all {model} vehicle records.")
            print(f"Aborted deletion of {model} vehicle records.")
            return
        delete_msg = {
            "data_type": "vehicle",
            "action": "delete",
            "port": self.port,
            "model": model
        }
        send_message("localhost", self.server_port, delete_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return

        success_msg = f"Successfully deleted all {model} vehicle records."
        print(success_msg)
        logging.info(success_msg)

    # Delete a laptop
    def delete_laptop(self):
        engin_name = input("Enter the name of the engineer this laptop is loaned to.\n(Leave empty if laptop is loaned to no engineer):")
        engin_name = engin_name.strip()
        delete_msg = {
            "data_type": "laptop",
            "action": "delete",
            "port": self.port,
            "engineer": engin_name
        }
        if engin_name == "":
            id = get_digit_choice(f"Enter the ID number of the un-loaned laptop to be deleted:",
                                   "Invalid Laptop ID. Enter a number > 0", 1, inf)
            delete_msg["id"] = id

        send_message("localhost", self.server_port, delete_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return
        
        success_msg = f"Successfully deleted laptop from the database"
        logging.info(success_msg)
        print(success_msg)


    # Delete contact details
    def delete_contact(self):
        engin_name = input("Enter the name of the engineer whose contact details should be deleted.\n(Leave empty to delete contacts by id):")
        engin_name = engin_name.strip()
        proceed = get_yes_no_choice(f"!!! WARNING !!! - All contact details for engineer {engin_name} will be deleted. Proceed? (Y/N):")
        if proceed == 'n':
            logging.info(f"User chose to abort deleting all contact details for engineer {engin_name}")
            print(f"Aborted deleting contact details for engineer {engin_name}.")
            return

        delete_msg = {
            "data_type": "contact_details",
            "action": "delete",
            "port": self.port,
            "engineer": engin_name
        }
        if engin_name == "":
            id = get_digit_choice(f"Enter the ID number of the contact details to be deleted:",
                                   "Invalid Contact Details ID. Enter a number > 0", 1, inf)
            delete_msg["id"] = id

        send_message("localhost", self.server_port, delete_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return
        
        success_msg = "Successfully deleted contact details."
        logging.info(success_msg)
        print(success_msg)


    # Decorator for query functions, prompts to dump the results as a JSON object/array
    def json_dump_prompt(func):
        def wrapper(*args, **kwargs):
            json_object = func(*args, **kwargs)
            if json_object is None:
                return None
            dump = get_yes_no_choice("Dump the result as JSON? (Y/N)")
            if dump == 'n':
                return json_object
            out_file_name = input("Enter the name of the file to dump JSON to:")
            logging.info(f"Attempting to dump user read result as JSON to file named {out_file_name}")
            dump_to_json(json_object, out_file_name)
            succcess_msg = f"Successfully dumped {json_object} to {out_file_name}"
            print(succcess_msg)
            logging.info(succcess_msg)
            return json_object
        return wrapper

    # Decorator for query functions, prompts to dump results to a Word document (docx)
    def docx_dump_prompt(func):
        def wrapper(*args, **kwargs):
            json_object = func(*args, **kwargs)
            if json_object is None:
                return None
            dump = get_yes_no_choice("Dump the result to a Word document? (Y/N):")
            if dump == 'n':
                return
            out_file_name = input("Enter the name of the word file to dump to:")
            # remove any potential file extension and replace with .docx
            out_file_name = out_file_name.rsplit('.', 1)[0] + ".docx"
            logging.info(f"Attempting to dump user read result to Word document named {out_file_name}")
            dump_to_docx(json_object, out_file_name)
            success_msg = f"Successfully dumped read result to Word document {out_file_name}"
            logging.info(success_msg)
            print(success_msg)
        return wrapper

    # Query vehicles, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_vehicles(self):
        model = input("Enter the model of vehicles to read.\nEnter \"all\" to read all vehicles.\n(Leave blank to read by vehicle ID):")
        model = model.strip()
        read_msg = {
            "data_type": "vehicle",
            "action": "read",
            "port": self.port,
            "model": model
        }
        if model == "":
            id = get_digit_choice("Enter the vehicle id to read:",
                                  "Invalid vehicle ID. Enter a number > 0", 1, inf)
            logging.info(f"Reading info for vehicle with ID {id}")
            read_msg["id"] = id
        else:
            logging.info(f"Asking the server to read info for vehicles of model {model}")

        send_message("localhost", self.server_port, read_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return None

        try:
            cars_json = server_response["vehicles"]
        except:
            error_msg = "Server response was marked as successful, but did not provide an entry \"vehicles\" containing data for vehicle information "
            logging.error(error_msg)
            print(error_msg)
            return None

        success_msg = "Successfully read vehicles. Vehicle info:"
        logging.info(success_msg)
        print(success_msg)

        for car in cars_json:
            logging.info(car)
            print(car)

        return cars_json

    # Query engineers, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_engineers(self):
        name = input("Enter the name of the engineer to read.\nEnter \"all\" to read all vehicles.\n(Leave blank to read by engineer ID):")
        name = name.strip()
        read_msg = {
            "data_type": "engineer",
            "action": "read",
            "port": self.port,
            "name": name
        }
        if name == "":
            id = get_digit_choice("Enter the engineer ID to read:",
                                  "Invalid engineer ID. Enter a number > 0", 1, inf)
            logging.info(f"Reading info for engineer with ID {id}")
            read_msg["id"] = id
        else:
            logging.info(f"Reading info for engineer named {name}")

        send_message("localhost", self.server_port, read_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return None

        try:
            engins_json = server_response["engineers"]
        except:
            error_msg = "Server response was marked successful, but did not include an entry \"engineers\" with data of the read engineers."
            logging.error(error_msg)
            print(error_msg)
            return None
        
        success_msg = "Successfully read engineers. Engineer info:"
        logging.info(success_msg)
        print(success_msg)

        for engin in engins_json:
            logging.info(engin)
            print(engin)
        
        return engins_json

    @docx_dump_prompt
    @json_dump_prompt
    # Query vehicle_engineers, either by engineer_id or vehicle_id
    def query_vehicle_engineers(self):
        read_msg = {
            "data_type": "vehicle_engineers",
            "action": "read",
            "port": self.port
        }
        read_vehicles = get_yes_no_choice("Read vehicles by engineer name? Enter \"N\" to read engineers by vehicle model. (Y/N):")
        if read_vehicles == 'y':
            name = input("Whose vehicles do you want to read? (Enter an engineer name):")
            logging.info(f"Asking server to read vehicles engineered by {name}")
            read_msg["engineer"] = name
        
        else:
            model = input("Which vehicle model's engineers do you want to read? (Enter a model name):")
            logging.info(f"Asking server to read engineers that worked on model {model} vehicles.")
            read_msg["model"] = model
        
        send_message("localhost", self.server_port, read_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return

        if read_vehicles == 'y':
            try:
                vehicles_json = server_response["vehicles"]
            except:
                error_msg = f"Client wanted to read vehicles made by engineer {name}, but server response did not include an entry for \"vehicles\""
                logging.error(error_msg)
                print(error_msg)
                return
            
            success_msg = f"Successfully read vehicles made by engineer {name}. Info for vehicles:"
            logging.info(success_msg)
            print(success_msg)

            for car_json in vehicles_json:
                logging.info(car_json)
                print(car_json)
            
            return vehicles_json
        
        else:
            try:
                engins_json = server_response["engineers"]
            except:
                error_msg = f"Client wanted to read engineers who worked on vehicle model {model}, but server response did not include an entry for \"engineers\""
                logging.error(error_msg)
                print(error_msg)
                return
            
            success_msg = f"Successfully read engineers who worked on vehicle model {model}. Info for engineers:"
            logging.info(success_msg)
            print(success_msg)

            for engin_json in engins_json:
                logging.info(engin_json)
                print(engin_json)
            
            return engins_json

    # Query laptops, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_laptops(self):
        read_all = get_yes_no_choice("Read all laptops? Enter \"N\" to read a single laptop (Y/N):")
        read_msg = {
            "data_type": "laptop",
            "action": "read",
            "port": self.port
        }
        if read_all == 'y':
            logging.info("Asking server to read info for all laptops")
            read_msg["model"] = "all"

        else:
            read_by_model = get_yes_no_choice("Read laptops by model name? Enter \"N\" to read by engineer name (Y/N):")
            if read_by_model == 'y':
                model = input("Enter the model of laptops to read:")
                logging.info(f"Asking server to read info for model {model} laptops")
                read_msg["model"] = model

            else:
                engin_name = input("Enter the name of engineer whose loaned laptop will be read:")
                logging.info(f"Asking server to read info for laptop loaned by engineer {engin_name}")
                read_msg["model"] = ""
                read_msg["engineer"] = engin_name

        send_message("localhost", self.server_port, read_msg)
        
        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return None

        try:
            laptops_json = server_response["laptops"]
        except:
            error_msg = "Server response for laptop read job was marked successful, but had no entry \"laptops\" with data for read laptops"
            logging.error(error_msg)
            print(error_msg)
            return None
        
        success_msg = "Successfully read laptops. Laptop info:"
        logging.info(success_msg)
        print(success_msg)

        for laptop in laptops_json:
            logging.info(laptop)
            print(laptop)
        
        return laptops_json

    # Query contact details, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_contacts(self):
        name = input("Enter the name of the engineer whose contact details will be read.\nEnter \"all\" to read all contact details.\nLeave blank to read contact details by ID:")
        name = name.strip()
        read_msg = {
            "data_type": "contact_details",
            "action": "read",
            "port": self.port,
            "engineer": name
        }
        
        if name == "":
            id = get_digit_choice("Enter the contact details ID to read:",
                                  "Invalid ID. Enter a number > 0", 1, inf)
            logging.info(f"Asking server to read info for contact details with ID {id}")
            read_msg["id"] = id

        else:
            logging.info(f"Asking server to read info for contact details of engineer {name}")
        
        send_message("localhost", self.server_port, read_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return None

        try:
            contacts_json = server_response["contact_details"]
        except:
            error_msg = "Server contact details read job had no entry \"contact_details\" with data for read contact details."
            logging.error(error_msg)
            print(error_msg)
            return None
        
        success_msg = "Successfully read contact details. Contact details info:"
        logging.info(success_msg)
        print(success_msg)

        for contact in contacts_json:
            logging.info(contact)
            print(contact)
        
        return contacts_json

    def update_vehicle(self):
        vehicle_id = get_digit_choice("Enter the ID of the vehicle to update:",
                                      "Invalid ID. Enter a number > 0", 1, inf)
        update_msg = {
            "data_type": "vehicle",
            "action": "update",
            "port": self.port,
            "id": vehicle_id
        }
        model = input(f"Enter the new model for vehicle ID {vehicle_id}\n(Leave blank to keep model the same):")
        if model != "":
            update_msg["model"] = model

        quantity = get_digit_choice(f"Enter the new quantity of vehicle ID {vehicle_id}\n(Enter \"-1\" to keep quantity the same):",
                                    "Please enter a non-negative quantity (or -1 to skip updating quantity).",
                                    -1, inf)
        if quantity > -1:
            update_msg["quantity"] = quantity

        price = get_digit_choice(f"Enter the new price of vehicle ID {vehicle_id}\n(Enter \"0\" to keep price the same):",
                                 "Invalid price. Please enter a price > 0 (or 0 to skip updating price)",
                                 0, inf)
        if price > 0:
            update_msg["price"] = price

        year = get_digit_choice(f"Enter the new year vehicle ID {vehicle_id} was manufactured\n(Enter \"0\" to keep year the same):",
                                "Invalid year. Please enter a year in the range (1920-2021) (or enter \"0\" to skip updating manufacture year",
                                0, 2022)
        if 1920 <= year < 2022:
            update_msg["manufacture_year"] = year
        elif 0 < year < 1920:
            logging.error(f"User entered invalid year {year}. Year should be between (1920-2021). Skipping update for vehicle year.")

        month = get_digit_choice(f"Enter the new month vehicle ID {vehicle_id} was manufactured\n(Enter \"0\" to keep month the same):",
                                "Invalid month. Please enter a month in the range (1-12) (or enter \"0\" to skip updating manufacture month",
                                0, 12)
        if month > 0:
            update_msg["manufacture_month"] = month

        day = get_digit_choice(f"Enter the new date vehicle ID {vehicle_id} was manufactured\n(Enter \"0\" to keep date the same):",
                                "Invalid date. Please enter a date in the range (1-32) (or enter \"0\" to skip updating manufacture date",
                                0, 32)
        if day > 0:
            update_msg["manufacture_date"] = day
        
        engineer_names = input(f"Enter comma-separated names (e.g. Engineer1, Engineer2, etc) of the updated engineers assigned to vehicle ID {vehicle_id}\n(Leave blank to keep assigned engineers the same):")
        engineer_names = engineer_names.split(",")
        engineer_names = [name.strip() for name in engineer_names]

        if engineer_names:
            update_msg["engineers"] = engineer_names
        
        send_message("localhost", self.server_port, update_msg)

        server_response = self.get_server_response()

        status = self.check_server_status(server_response)

        if not status:
            return

        try:
            vehicle_json = server_response["vehicle"]
        except:
            error_msg = f"Update job for vehicle ID {vehicle_id} was marked successful, but the server did not provide an entry for \"vehicle\""
            logging.error(error_msg)
            print(error_msg)
            return None

        success_msg = f"Successfully updated info for vehicle ID {vehicle_id}. Vehicle new info:\n{vehicle_json}"

        logging.info(success_msg)
        print(success_msg)

        return vehicle_json

        
    def update_engineer(self):
        print("called update_engineer")

    def update_laptop(self):
        print("called update_laptop")

    def update_contacts(self):
        print("called update_contacts")

    def insert_from_json(self):
        file_name = input("Enter the file name containing JSON data to insert:")
        file_name = file_name.strip()
        try:
            logging.info(f"Attempting to insert JSON data stored in file {file_name} to the database")
            with open(file_name) as f:
                json_data = json.load(f)
            if isinstance(json_data, dict):
                self.parse_json_object_and_insert(json_data)
            else:
                # JSON array
                for json_object in json_data:
                    self.parse_json_object_and_insert(json_object)
        except JSONDecodeError:
            error_msg = f"JSONDecodeError: There was a problem reading the JSON object in {file_name} ; JSON object likely not encoded properly."
            print(error_msg)
            logging.error(error_msg)
        except FileNotFoundError:
            error_msg = f"FileNotFoundError: file named {file_name} does not exist."
            print(error_msg)
            logging.error(error_msg)

    def parse_json_object_and_insert(self, json_object):
        logging.info(f"Attempting to parse JSON object and insert into the database.")
        
        data_type = json_object.get('data_type', None)
        insert_msg = {
            "data_type": data_type,
            "action": "add",
            "port": self.port
        }
        if data_type is None:
            error_msg = "Error: no \"data_type\" entry found in JSON object." + \
                        f"\nAborted adding {json_object} to the database."
            print(error_msg)
            logging.error(error_msg)
        
        send_message("localhost", self.server_port, {**insert_msg, **json_object})

        server_response = self.get_server_response()
        status = self.check_server_status(server_response)
        if not status:
            return

        elif data_type == "vehicle":
            logging.info("Attempting to add a vehicle to the database from JSON object.")
            
            try:
                model = server_response["model"]
                quantity = server_response["quantity"]
                price = server_response["price"]
                manufacture_year = server_response["manufacture_year"]
                manufacture_month = server_response["manufacture_month"]
                manufacture_date = server_response["manufacture_date"]
            except:
                error_msg = "Server JSON vehicle insert job is missing one of [\"model\", \"quantity\", \"price\", \"manufacture_year\", \"manufacture_month\", \"manufacture_date\"]"
                logging.error(error_msg)
                print(error_msg)
                return

            success_msg = "New vehicle info:" + "\n" + \
                          f"Model: {model}\nQuantity: {quantity}\nPrice: {price}\nManufacture Year: {manufacture_year}\nManufacture Month: {manufacture_month}\nManufacture Date: {manufacture_date}" + \
                          f"\nSuccessfully added new vehicle to the database!"
            print(success_msg)
            logging.info(success_msg)

        elif data_type == "engineer":
            logging.info("Attempting to add an engineer to the database from JSON object.")

            try:
                name = server_response["name"]
                birth_year = server_response["birth_year"]
                birth_month = server_response["birth_month"]
                birth_date = server_response["birth_date"]
            except:
                error_msg = "Server JSON engineer insert job is missing one of [\"name\", \"birth_year\", \"birth_month\", \"birth_date\"]"
                print(error_msg)
                logging.error(error_msg)
                return

            success_msg = "New engineer info:\n" + f"Name: {name}\nBirth Year: {birth_year}\nBirth Month: {birth_month}\nBirth Date: {birth_date}" + \
                          f"\nSuccessfully added new engineer to the database!"
            print(success_msg)
            logging.info(success_msg)
                

        elif data_type == "laptop":
            logging.info("Attempting to add a laptop to the database from JSON object")
            try:
                model = json_object.get('model', None)
                loan_year = json_object.get('loan_year', None)
                loan_month = json_object.get('loan_month', None)
                loan_date = json_object.get('loan_date', None)
                engineer = json_object.get('engineer', None)
            except:
                error_msg = "Server JSON laptop insert job is missing one of [\"model\", \"loan_year\", \"loan_month\", \"loan_date\", \"engineer\"]"
                print(error_msg)
                logging.error(error_msg)
                return
                
            success_msg = "New laptop info:\n" + f"Model: {model}\nLoan Year: {loan_year}\nLoan Month: {loan_month}" + \
                          "\nLoan Date: {loan_date}\nEngineer: {engineer}" + \
                          "\nSuccessfully added new laptop to the database!"
            print(success_msg)
            logging.info(success_msg)

        elif data_type == "contact_details":
            logging.info("Attempting to add new contact details from JSON object.")
            try:
                phone_number = json_object.get('phone_number', None)
                address = json_object.get('address', None)
                engineer = json_object.get('engineer', None)
            except:
                error_msg = "Server JSON contact details insert job is missing one of [\"phone_number\", \"address\", \"engineer\"]"
                print(error_msg)
                logging.error(error_msg)
                return
            
            success_msg = "New contact details info:\n" + \
                          f"Phone Number: {phone_number}\nAddress: {address}\nEngineer: {engineer}" + \
                          "\nSuccessfully added new contact details to the database!"
            print(success_msg)
            logging.info(success_msg)

        else:
            data_type_error = f"\"data_type\" entry value of {data_type} is not recognized.\n" + \
                              "\"data_type\" should be one of: \"vehicle\", \"engineer\", \"laptop\", \"contact_details\", \"vehicle_engineers\"\n" + \
                              f"Aborted adding entry {json_object} to the database."
            print(data_type_error)
            logging.error(data_type_error)
            return

    def display_interface(self):
        tables = ["Vehicles", "Engineers", "Laptops", "Contact Details", "Vehicle-Engineer Assignments"]
        function_map = [
            [self.add_vehicle, self.delete_vehicle, self.query_vehicles, self.update_vehicle],
            [self.add_engineer, self.delete_engineer, self.query_engineers, self.update_engineer],
            [self.add_laptop, self.delete_laptop, self.query_laptops, self.update_laptop],
            [self.add_contact, self.delete_contact, self.query_contacts, self.update_contacts],
            self.insert_from_json
        ]
        print("\nWelcome to your Database Utilities!")
        print("1. Vehicles")
        print("2. Engineers")
        print("3. Laptops")
        print("4. Contact Details")
        print("5. Vehicle-Engineer Assignments (Read-Only)")
        print("6. Insert data from a JSON file")
        print("7: Exit Database Utilities")

        table_choice = get_digit_choice("Select a table, JSON insertion, or Exit:",
                                        "Invalid selection. Please enter a digit corresponding to the desired table, JSON insert, or Exit.",
                                        1, 8)
        table_choice -= 1

        if table_choice == 6:
            return False

        if table_choice == 5:
            self.insert_from_json()
            return True
        
        if table_choice == 4:
            self.query_vehicle_engineers()
            return True

        print(f"\nWhat would you like to do in the {tables[table_choice]} table?")
        print(f"1. Add a(n) {tables[table_choice][0:-1]}")
        print(f"2. Delete a(n) {tables[table_choice][0:-1]}")
        print(f"3. Read from the {tables[table_choice]} table")
        print(f"4. Update a(n) {tables[table_choice][0:-1]}")
        print("5. Choose a different table to work with")

        action_choice = get_digit_choice(f"Select an action to perform on the {tables[table_choice]} table:",
                                         "Invalid selection. Please enter a digit corresponding to the desired action.",
                                         1, 6)
        action_choice -= 1

        if action_choice == 4:
            # tell main() to restart the interface
            return True

        # call the chosen function from the function map
        function_map[table_choice][action_choice]()

        exit_choice = get_yes_no_choice("Exit the utilities script? (Y/N):")

        # returns false if user chose y to exit, tells main to stop running interface()
        return exit_choice.lower() != "y"

@click.command()
@click.argument("server_port", nargs=1, type=int)
@click.argument("client_port", nargs=1, type=int)
def main(server_port, client_port):
    print(f"Server port: {server_port}")
    print(f"Client port: {client_port}")
    client = Client(server_port, client_port)
    run_again = client.display_interface()
    while run_again:
        run_again = client.display_interface()

if __name__ == "__main__":
    main()