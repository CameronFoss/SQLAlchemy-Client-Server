from json.decoder import JSONDecodeError
from math import inf
from datetime import date

from sqlalchemy.orm.exc import UnmappedInstanceError
from training.choice_funcs import get_digit_choice, get_yes_no_choice, get_day, get_month, get_year
from training.db_utils import VehicleUtils, EngineerUtils, LaptopUtils, ContactDetailsUtils, cleanup_utils
from docx import Document
import json
import logging

def dump_to_json(object, out_file_name):
    try:
        with open(out_file_name, 'w') as f:
            json.dump(object, f, indent=4, sort_keys=True)
    except JSONDecodeError:
        msg = f"JSONDecodeError: There was a problem dumping the JSON object {object} to {out_file_name}."
        print(msg)
        logging.error(msg)

def dump_to_docx(object, out_file_name):
    document = Document()
    document.add_heading("Query Results", 0)
    if isinstance(object, dict):
        # only have one row worth of data
        table = document.add_table(rows=1, cols=len(object))
        table.autofit = True
        # iterate over keys for header cells and vals for data cells
        hdr_cells = table.rows[0].cells
        row_cells = table.add_row().cells
        for i, key, value in enumerate(object.items()):
            hdr_cells[i].text = key
            row_cells[i].text = str(value)
    else:
        # iterate over object for many rows of data
        table = document.add_table(rows=1, cols=len(object[0]))
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(object[0]):
            hdr_cells[i].text = key
        for dic in object:
            row_cells = table.add_row().cells
            for i, value in enumerate(dic.values()):
                row_cells[i].text = str(value)
    document.add_page_break()
    document.save(out_file_name)


class Interface:

    def __init__(self):
        self.car_utils = VehicleUtils()
        self.engin_utils = EngineerUtils()
        self.laptop_utils = LaptopUtils()
        self.contact_utils = ContactDetailsUtils()
        logging.basicConfig(filename="utilities.log", level=logging.DEBUG, format="%(asctime)s - %(levelname)s: %(message)s")
    
    def __del__(self):
        cleanup_utils()

    # Add a vehicle to the DB, update quantity if model already exists
    def add_vehicle(self, model=None, prompt_engin_assign=True):
        if model is None:
            model = input("Enter the model of the new vehicle:")
        quantity = get_digit_choice("Enter the quantity of new vehicles:",
                                    "Please enter a non-negative quantity.",
                                    0, inf)
        price = get_digit_choice("Enter the price of the new vehicle:",
                                 "Invalid price. Please enter a price > 0",
                                 1, inf)
        year = get_year("Enter the year the new vehicle was manufactured:")
        month = get_month("Enter the month the new vehicle was manufactured:")
        day = get_day("Enter the date the new vehicle was manufactured:")
        manufacture_date = date(year, month, day)
        new_car = self.car_utils.add_vehicle_db(model, quantity, price, manufacture_date)
        logging.info(f"Attempting to add vehicle {new_car} to the database.")
        if new_car is None:
            # Updated quantity, no more work to do
            msg1 = f"Vehicle model {model} manufactured on {manufacture_date} already exists in the database."
            msg2 = f"Increased quantity of vehicle model {model} by {quantity}."
            print(msg1)
            print(msg2)
            logging.info(msg1)
            logging.info(msg2)
            return

        # Added a new vehicle, prompt user to assign engineers to the vehicle
        success_msg = f"Successfully added new vehicle!\nModel: {model}\nQuantity: {quantity}\nPrice: {price}\nManufacture Date: {manufacture_date}"
        print(success_msg)
        logging.info(success_msg)

        if prompt_engin_assign:
            assign_engin = get_yes_no_choice("Assign engineers to this vehicle? (Y/N):")
            if assign_engin == 'y':
                engineer_names = input("\nEnter names of engineers to assign.\nSeparate names by a comma (e.g. John Smith, Jane Doe):")
                engineer_names = engineer_names.split(',')
                new_engins = []
                logging.info(f"Attempting to assign engineers {engineer_names} to the {model} manufactured on {manufacture_date}.")
                for name in engineer_names:
                    name = name.strip()
                    engin = self.engin_utils.read_engineer_by_name(name)
                    if engin is None:
                        logging.info(f"Engineer {name} didn't exist in the database.")
                        add_engin = get_yes_no_choice(f"Engineer {name} does not exist in the database.\nWould you like to add them? (Y/N):")
                        if add_engin == 'y':
                            logging.info(f"Adding engineer {name} to the database before assigning them to the vehicle.")
                            engin = self.add_engineer(engin_name=name, prompt_vehicle_assign=False)
                        else:
                            # skip trying to add this engineer below if user doesn't want to add them to the DB
                            logging.info(f"User skipped adding engineer {name} to the database.")
                            continue
                    new_engins += [engin]
                    self.car_utils.update_vehicle_db(new_car.id, engineers=new_engins)
                    print(f"Engineer {name} assigned to the new vehicle.")
                    logging.info(f"Engineer {name} assigned to the {model} manufactured on {manufacture_date}")
                print("\nSuccessfully assigned engineers to the new vehicle!")
                logging.info(f"\nSuccessfully assigned engineers to vehicle model {model} manufactured on {manufacture_date}")
            else:
                logging.info(f"User skipped assigning any engineers to new vehicle {model} manufactured on {manufacture_date}")
        return new_car

    # Add an engineer to the DB
    def add_engineer(self, engin_name=None, prompt_vehicle_assign=True):
        if engin_name is None:
            engin_name = input("Enter the new Engineer's name:")
        birth_year = get_year(f"Enter engineer {engin_name}'s birth year:")
        birth_month = get_month(f"Enter engineer {engin_name}'s birth month:")
        birth_day = get_day(f"Enter engineer {engin_name}'s birth date:")
        date_of_birth = date(birth_year, birth_month, birth_day)
        new_engin = self.engin_utils.add_engineer_db(engin_name, date_of_birth)
        success_msg = f"Successfully added new engineer {engin_name} born on {date_of_birth}"
        print(success_msg)
        logging.info(success_msg)
        
        if prompt_vehicle_assign:
            assign_vehicle = get_yes_no_choice("Assign the new engineer to any vehicles? (Y/N):")
            if assign_vehicle == 'y':
                vehicle_models = input("Enter models of vehicles to assign.\nIt is assumed that manufacture date is irrelevant.\nSeparate model names by a comma (e.g. Fusion, Bronco):")
                vehicle_models = vehicle_models.split(',')
                logging.info(f"Attempting to assign engineer {engin_name} to vehicles {vehicle_models}")
                for model in vehicle_models:
                    model = model.strip()
                    cars = self.car_utils.read_vehicles_by_model(model)
                    if not cars:
                        logging.info(f"No vehicles of model {model} existed in the database.")
                        add_car = get_yes_no_choice(f"No vehicles of model {model} exist in the database. Would you like to add one? (Y/N):")
                        if add_car == 'y':
                            logging.info(f"Attempting to add vehicle model {model} to the database.")
                            cars.append(self.add_vehicle(model, prompt_engin_assign=False))
                        else:
                            # Skip trying to add this model below and move to the next one
                            logging.info(f"User skipped adding non-existant vehicle model {model} to the database")
                            continue
                    for car in cars:
                        new_engins_list = car.engineers + [new_engin]
                        self.car_utils.update_vehicle_db(car.id, engineers=new_engins_list)
                        assignment_msg = f"New engineer {engin_name} successfully assigned to model {car.model} manufactured on {car.manufacture_date}"
                        print(assignment_msg)
                        logging.info(assignment_msg)
                print(f"Successfully assigned {engin_name} to vehicles!")
                logging.info(f"Successfully assigned engineer {engin_name} to vehicles.")
            else:
                logging.info(f"User skipped assigning engineer {engin_name} to any vehicles.")
        
        return new_engin

    # Add a laptop to the DB and loan to an engineer if desired
    def add_laptop(self):
        engin_name = input("Enter the name of the engineer this laptop will be loaned to:")
        engin_name = engin_name.strip()
        logging.info(f"Adding a new laptop to the database. Attempting to loan it to engineer {engin_name}")
        engin = self.engin_utils.read_engineer_by_name(engin_name)
        if engin is None:
            logging.info(f"Engineer {engin_name} did not exist in the database.")
            add_engin = get_yes_no_choice(f"Engineer {engin_name} does not exist in the database. Would you like to add them? (Y/N):")
            if add_engin == 'y':
                logging.info(f"Attempting to add engineer {engin_name} to the database.")
                engin = self.add_engineer(engin_name)
            else:
                abort_laptop = get_yes_no_choice(f"Would you still like to add this laptop without loaning it to an engineer? (Y/N):")
                if abort_laptop == 'n':
                    logging.info(f"User chose to abort adding a new laptop since engineer {engin_name} did not exist in the database.")
                    print("Aborted adding new laptop to the database.")
                    return None
        # if engineer already has a laptop, prompt to replace it with the new one
        prev_laptop = self.laptop_utils.read_laptop_by_owner(engin_name)
        if prev_laptop is not None:
            warning_msg = f"Engineer {engin_name} already has a laptop loaned to them\n" + \
                          f"\n!!! WARNING !!! - Adding a new laptop would replace the laptop already loaned to {engin_name}" + \
                          f"{engin_name}'s previous laptop would not be deleted from the database, but would be loaned by no one."
            print(warning_msg)
            logging.warning(warning_msg)
            replace_laptop = get_yes_no_choice(f"Replace {engin_name}'s laptop with the new one? (Y/N)")
            if replace_laptop == 'n':
                logging.info(f"User aborted adding a new laptop to the database as to not replace engineer {engin_name}'s current laptop.")
                print(f"Aborted adding new laptop to the database.")
                return
        model = input("Enter the model of the new laptop:")
        year_loaned = get_year(f"Enter the year the laptop was loaned to {engin_name}:")
        month_loaned = get_month(f"Enter the month the laptop was loaned to {engin_name}:")
        day_loaned = get_day(f"Enter the date the laptop was loaned to {engin_name}")
        date_loaned = date(year_loaned, month_loaned, day_loaned)
        new_laptop = self.laptop_utils.add_laptop_db(model, date_loaned, engin_name)
        if engin is None:
            no_engin_msg = f"Successfully added new {model} laptop, but it is not loaned by any engineer."
            print(no_engin_msg)
            logging.info(no_engin_msg)
        else:
            loaned_msg = f"Successfully loaned new {model} laptop to {engin_name}"
            print(loaned_msg)
            logging.info(loaned_msg)
        return new_laptop

    # Add contact details for an engineer
    def add_contact(self):
        engin_name = input("Which engineer's contact info are you providing? Enter their name:")
        engin_name = engin_name.strip()
        engin = self.engin_utils.read_engineer_by_name(engin_name)
        logging.info(f"Attempting to add contact details for engineer {engin_name}.")
        if engin is None:
            doesnt_exist_msg = f"Engineer {engin_name} does not exist in the database."
            print(doesnt_exist_msg)
            logging.info(doesnt_exist_msg)
            add_engin = get_yes_no_choice(f"Contact details cannot be added without an existing engineer. Would you like to add {engin_name} as a new engineer? (Y/N):")
            if add_engin == 'n':
                logging.info(f"User chose to abort adding contact details for non-existant engineer {engin_name}.")
                print(f"Aborting adding contact details for non-existant engineer {engin_name}.")
                return None
            logging.info(f"Attempting to add engineer {engin_name} to the database before adding contact details.")
            engin = self.add_engineer(engin_name)
        phone_number = input(f"Enter {engin_name}'s phone number (XXX)-XXX-XXXX:")
        address = input(f"Enter {engin_name}'s address:")
        new_contact = self.contact_utils.add_contact_details_db(phone_number, address, engin_name)
        success_msg = f"Successfully added new contact information for {engin_name}!"
        print(success_msg)
        logging.info(success_msg)
        return new_contact
        
    # Delete an engineer
    def delete_engineer(self):
        engin_name = input("Enter the name of the engineer to be deleted:")
        engin_name = engin_name.strip()
        logging.info(f"Attempting to delete engineer {engin_name} from the database.")
        engin = self.engin_utils.read_engineer_by_name(engin_name)
        if engin is None:
            doesnt_exist_msg = f"Engineer {engin_name} does not exist in the database." + \
                               f"Aborting deleting non-existant engineer {engin_name}."
            print(doesnt_exist_msg)
            logging.info(doesnt_exist_msg)
            return
        proceed = get_yes_no_choice(f"!!! WARNING !!! - Deleting engineer {engin_name} will also delete any of their contact details. Proceed? (Y/N):")
        if proceed == 'n':
            print(f"Aborted deleting engineer {engin_name} from the database.")
            logging.info(f"User chose to abort deleting engineer {engin_name} as to not delete their contact details.")
            return
        self.engin_utils.delete_engineer_by_name(engin_name)
        success_msg = f"Successfully deleted engineer {engin_name} and their contact details from the database."
        print(success_msg)
        logging.info(success_msg)

    # Delete a vehicle
    def delete_vehicle(self):
        model = input("Enter the model of the vehicle to delete:")
        model = model.strip()
        logging.info(f"Attempting to delete all {model} vehicle records from the database.")
        proceed = get_yes_no_choice(f"!!! WARNING !!! - ALL {model} vehicle records will be deleted, regardless of manufacture date.\nProceed? (Y/N):")
        if proceed == 'n':
            logging.info(f"User chose to abort deleting all {model} vehicle records.")
            print(f"Aborted deletion of {model} vehicle records.")
            return
        self.car_utils.delete_vehicle_by_model(model)
        success_msg = f"Successfully deleted all {model} vehicle records."
        print(success_msg)
        logging.info(success_msg)

    # Delete a laptop
    def delete_laptop(self):
        engin_name = input("Enter the name of the engineer this laptop is loaned to.\n(Leave empty if laptop is loaned to no engineer):")
        engin_name = engin_name.strip()
        if engin_name == "":
            id = get_digit_choice(f"Enter the ID number of the un-loaned laptop to be deleted:",
                                   "Invalid Laptop ID. Enter a number > 0", 1, inf)
            try:
                logging.info(f"Attempting to delete laptop with ID {id}")
                self.laptop_utils.delete_laptop_by_id(id)
                success_msg = f"Successfully deleted laptop with ID {id}!"
                print(success_msg)
                logging.info(success_msg)
            except UnmappedInstanceError:
                error_msg = f"Laptop with ID {id} has already been deleted from the database."
                print(error_msg)
                logging.error(error_msg)
                return
        else:
            logging.info(f"Attempting to delete laptop loaned to engineer {engin_name} from the database.")
            self.laptop_utils.delete_laptop_by_owner(engin_name)
            success_msg = f"Successfully deleted laptop loaned by {engin_name}!"
            print(success_msg)
            logging.info(success_msg)

    # Delete contact details
    def delete_contact(self):
        engin_name = input("Enter the name of the engineer whose contact details should be deleted.\n(Leave empty to delete contacts by id):")
        engin_name = engin_name.strip()
        proceed = get_yes_no_choice(f"!!! WARNING !!! - All contact details for engineer {engin_name} will be deleted. Proceed? (Y/N):")
        if proceed == 'n':
            logging.info(f"User chose to abort deleting all contact details for engineer {engin_name}")
            print(f"Aborted deleting contact details for engineer {engin_name}.")
            return
        if engin_name == "":
            id = get_digit_choice(f"Enter the ID number of the contact details to be deleted:",
                                   "Invalid Contact Details ID. Enter a number > 0", 1, inf)
            try:
                logging.info(f"Attempting to delete contact details with ID {id}")
                self.contact_utils.delete_contact_details_by_id(id)
            except UnmappedInstanceError:
                error_msg = f"Contact Details with id {id} has already been deleted from the database."
                print(error_msg)
                logging.error(error_msg)
                return
        else:
            logging.info(f"Attempting to delete all contact details for engineer {engin_name}")
            engin = self.engin_utils.read_engineer_by_name(engin_name)
            if engin is None:
                doesnt_exist_msg = f"Engineer {engin_name} does not exist in the database." + \
                                   f"Aborted deleting contact details for non-existant engineer {engin_name}."
                print(doesnt_exist_msg)
                logging.info(doesnt_exist_msg)
                return
            try:
                self.contact_utils.delete_contact_details_by_engin_id(engin.id)
            except UnmappedInstanceError:
                error_msg = f"Engineer {engin_name} has no contact details to delete." + \
                            f"Aborted deleting contact details for engineer {engin_name}."
                print(error_msg)
                logging.error(error_msg)

    # Decorator for query functions, prompts to dump the results as a JSON object/array
    def json_dump_prompt(func):
        def wrapper(*args, **kwargs):
            db_object = func(*args, **kwargs)
            dump = get_yes_no_choice("Dump the result as JSON? (Y/N)")
            if dump == 'n':
                return db_object
            out_file_name = input("Enter the name of the file to dump JSON to:")
            logging.info(f"Attempting to dump user read result as JSON to file named {out_file_name}")
            if isinstance(db_object, dict):
                # db_object is a single json object
                json_obj = db_object.to_json()
                dump_to_json(json_obj, out_file_name)
                succcess_msg = f"Successfully dumped {json_obj} to {out_file_name}"
                print(succcess_msg)
                logging.info(succcess_msg)
                return db_object
            else:
                # db_object is a list of json objects
                json_array = [obj.to_json() for obj in db_object]
                dump_to_json(json_array, out_file_name)
                succcess_msg = f"Successfully dumped {json_array} to {out_file_name}"
                print(succcess_msg)
                logging.info(succcess_msg)
                return db_object
        return wrapper

    # Decorator for query functions, prompts to dump results to a Word document (docx)
    def docx_dump_prompt(func):
        def wrapper(*args, **kwargs):
            db_object = func(*args, **kwargs)
            dump = get_yes_no_choice("Dump the result to a Word document? (Y/N):")
            if dump == 'n':
                return
            out_file_name = input("Enter the name of the word file to dump to:")
            # remove any potential file extension and replace with .docx
            out_file_name = out_file_name.rsplit('.', 1)[0] + ".docx"
            logging.info(f"Attempting to dump user read result to Word document named {out_file_name}")
            if isinstance(db_object, dict):
                # db_object is a single DB object
                json_obj = db_object.to_json()
                dump_to_docx(json_obj, out_file_name)
            else:
                # db_object is a list of DB objects
                json_array = [obj.to_json() for obj in db_object]
                dump_to_docx(json_array, out_file_name)
        return wrapper

    # Query vehicles, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_vehicles(self):
        model = input("Enter the model of vehicles to read.\nEnter \"all\" to read all vehicles.\n(Leave blank to read by vehicle ID):")
        model = model.strip()
        if model == "":
            id = get_digit_choice("Enter the vehicle id to read:",
                                  "Invalid vehicle ID. Enter a number > 0", 1, inf)
            logging.info(f"Reading info for vehicle with ID {id}")
            car = self.car_utils.read_vehicle_by_id(id)
            print(f"Info for vehicle id {id}:")
            print(str(car))
            return car
        elif model == "all":
            logging.info("Reading all vehicle info.")
            cars = self.car_utils.read_vehicles_all()
            print("Info for all vehicles:")
            for car in cars:
                print(str(car))
            return cars
        else:
            logging.info(f"Reading info for model {model} vehicles.")
            cars = self.car_utils.read_vehicles_by_model(model)
            print(f"Info for vehicles of model {model}:")
            for car in cars:
                print(str(car))
            return cars

    # Query engineers, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_engineers(self):
        name = input("Enter the name of the engineer to read.\nEnter \"all\" to read all vehicles.\n(Leave blank to read by engineer ID):")
        name = name.strip()
        if name == "":
            id = get_digit_choice("Enter the engineer ID to read:",
                                  "Invalid engineer ID. Enter a number > 0", 1, inf)
            logging.info(f"Reading info for engineer with ID {id}")
            engin = self.engin_utils.read_engineer_by_id(id)
            print(f"Info for engineer ID {id}:")
            print(str(engin))
            return engin
        elif name == "all":
            engins = self.engin_utils.read_all_engineers()
            logging.info("Reading info for all engineers.")
            print("Info for all engineers:")
            for engin in engins:
                print(str(engin))
            return engins
        else:
            logging.info(f"Reading info for engineer named {name}")
            engin = self.engin_utils.read_engineer_by_name(name)
            print(f"Info for engineer {name}:")
            print(str(engin))
            return engin

    # Query laptops, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_laptops(self):
        read_all = get_yes_no_choice("Read all laptops? Enter \"N\" to read a single laptop (Y/N):")
        if read_all == 'y':
            logging.info("Reading info for all laptops")
            laptops = self.laptop_utils.read_all_laptops()
            print("Info for all laptops:")
            for laptop in laptops:
                print(str(laptop))
            return laptops
        read_by_model = get_yes_no_choice("Read laptops by model name? Enter \"N\" to read by engineer name (Y/N):")
        if read_by_model == 'y':
            model = input("Enter the model of laptops to read:")
            logging.info(f"Reading info for model {model} laptops")
            laptops = self.laptop_utils.read_laptops_by_model(model)
            print(f"Info for laptops of model {model}:")
            for laptop in laptops:
                print(str(laptop))
            return laptops
        else:
            engin_name = input("Enter the name of engineer whose loaned laptop will be read:")
            logging.info(f"Reading info for laptop loaned by engineer {engin_name}")
            laptop = self.laptop_utils.read_laptop_by_owner(engin_name)
            print(f"Info for laptop loaned by {engin_name}:")
            print(str(laptop))
            return laptop

    # Query contact details, prompt for each column
    @docx_dump_prompt
    @json_dump_prompt
    def query_contacts(self):
        name = input("Enter the name of the engineer whose contact details will be read.\nEnter \"all\" to read all contact details.\nLeave blank to read contact details by ID:")
        name = name.strip()
        if name == "all":
            logging.info("Reading info for all contact details")
            contacts = self.contact_utils.read_all_contact_details()
            print("Info for all contact details:")
            for contact in contacts:
                print(str(contact))
            return contacts
        elif name == "":
            id = get_digit_choice("Enter the contact details ID to read:",
                                  "Invalid ID. Enter a number > 0", 1, inf)
            logging.info(f"Reading info for contact details with ID {id}")
            contact = self.contact_utils.read_contact_details_by_id(id)
            print(f"Info for contact details with ID {id}:")
            print(str(contact))
            return contact
        else:
            logging.info(f"Reading info for contact details of engineer {name}")
            engin = self.engin_utils.read_engineer_by_name(name)
            contacts = self.contact_utils.read_contact_details_by_engin_id(engin.id)
            print(f"Info for contact details of engineer {name}:")
            for contact in contacts:
                print(str(contact))
            return contacts

    def insert_from_json(self):
        file_name = input("Enter the file name containing JSON data to insert:")
        file_name = file_name.strip()
        try:
            logging.info(f"Attempting to insert JSON data stored in file {file_name} to the database")
            with open(file_name) as f:
                json_data = json.load(f)
            if isinstance(json_data, dict):
                self.parse_json_object_and_insert(json_data)
            else:
                # JSON array
                for json_object in json_data:
                    self.parse_json_object_and_insert(json_object)
        except JSONDecodeError:
            error_msg = f"JSONDecodeError: There was a problem reading the JSON object in {file_name} ; JSON object likely not encoded properly."
            print(error_msg)
            logging.error(error_msg)
        except FileNotFoundError:
            error_msg = f"FileNotFoundError: file named {file_name} does not exist."
            print(error_msg)
            logging.error(error_msg)

    def parse_json_object_and_insert(self, json_object):
        logging.info(f"Attempting to parse JSON object and insert into the database.")
        data_type = json_object.get('data_type', None)
        if data_type is None:
            error_msg = "Error: no \"data_type\" entry found in JSON object." + \
                        f"\nAborted adding {json_object} to the database."
            print(error_msg)
            logging.error(error_msg)

        elif data_type == "vehicle":
            logging.info("Attempting to add a vehicle to the database from JSON object.")
            model = json_object.get('model', None)
            quantity = json_object.get('quantity', None)
            price = json_object.get('price', None)
            manufacture_year = json_object.get('manufacture_year', None)
            manufacture_month = json_object.get('manufacture_month', None)
            manufacture_date = json_object.get('manufacture_date', None)
            abort = False
            if model is None:
                model_error = "Error: no entry for vehicle \"model\" found."
                print(model_error)
                logging.error(model_error)
                abort = True
            if quantity is None:
                quantity_error = "Error: no entry for vehicle \"quantity\" found."
                print(quantity_error)
                logging.error(quantity_error)
                abort = True
            if price is None:
                price_error = "Error: no entry for vehicle \"price\" found."
                print(price_error)
                logging.error(price_error)
                abort = True
            if manufacture_year is None:
                year_error = "Error: no entry for vehicle \"manufacture_year\" found."
                print(year_error)
                logging.error(year_error)
                abort = True
            if manufacture_month is None:
                month_error = "Error: no entry for vehicle \"manufacture_month\" found."
                print(month_error)
                logging.error(month_error)
                abort = True
            if manufacture_date is None:
                date_error = "Error: no entry for vehicle \"manufacture_date\" found."
                print(date_error)
                logging.error(date_error)
                abort = True
            if abort:
                abort_msg = f"Aborted adding vehicle {json_object} to the database."
                print(abort_msg)
                logging.error(abort_msg)
                return
            new_car = self.car_utils.add_vehicle_db(model, quantity, price, date(manufacture_year, manufacture_month, manufacture_date))
            if new_car is None:
                car_exists = f"Vehicle model {model} manufactured on {date(manufacture_year, manufacture_month, manufacture_date)} " + \
                             f"already exists in the database.\nIncreased quantity of vehicle model {model} by {quantity}"
                print(car_exists)
                logging.info(car_exists)
            else:
                success_msg = "New vehicle info:" + "\n" + str(new_car) + f"\nSuccessfully added new vehicle to the database!"
                print(success_msg)
                logging.info(success_msg)

        elif data_type == "engineer":
            logging.info("Attempting to add an engineer to the database from JSON object.")
            name = json_object.get('name', None)
            birth_year = json_object.get('birth_year', None)
            birth_month = json_object.get('birth_month', None)
            birth_date = json_object.get('birth_date', None)
            abort = False
            if name is None:
                name_error = "Error: no entry for engineer \"name\" found."
                print(name_error)
                logging.error(name_error)
                abort = True
            if birth_year is None:
                year_error = "Error: no entry for engineer \"birth_year\" found."
                print(year_error)
                logging.error(year_error)
                abort = True
            if birth_month is None:
                month_error = "Error: no entry for engineer \"birth_month\" found."
                print(month_error)
                logging.error(month_error)
                abort = True
            if birth_date is None:
                date_error = "Error: no entry for engineer \"birth_date\" found."
                print(date_error)
                logging.error(date_error)
                abort = True
            if abort:
                abort_msg = f"Aborted adding engineer {json_object} to the database"
                print(abort_msg)
                logging.error(abort_msg)
                return
            new_engin = self.engin_utils.add_engineer_db(name, date(birth_year, birth_month, birth_date))
            if new_engin is None:
                dup_error = f"Engineer named {name} already exists in the database." + \
                            f"\nUpdating info for engineer {name} in the database."
                print(dup_error)
                logging.info(dup_error)
                engin = self.engin_utils.read_engineer_by_name(name)
                engin = self.engin_utils.update_engineer_by_id(engin.id, name=name, date_of_birth=date(birth_year, birth_month, birth_date))
                success_msg = f"Successfully updated info for engineer {name}:\n" + str(engin)
                print(success_msg)
                logging.info(success_msg)
                return
            success_msg = "New engineer info:\n" + str(new_engin) + f"\nSuccessfully added new engineer to the database!"
            print(success_msg)
            logging.info(success_msg)
                

        elif data_type == "laptop":
            logging.info("Attempting to add a laptop to the database from JSON object")
            model = json_object.get('model', None)
            loan_year = json_object.get('loan_year', None)
            loan_month = json_object.get('loan_month', None)
            loan_date = json_object.get('loan_date', None)
            engineer = json_object.get('engineer', None)
            abort = False
            if model is None:
                model_error = "Error: no entry for laptop \"model\" found."
                print(model_error)
                logging.error(model_error)
                abort = True
            if loan_year is None:
                year_error = "Error: no entry for laptop \"loan_year\" found."
                print(year_error)
                logging.error(year_error)
                abort = True
            if loan_month is None:
                month_error = "Error: no entry for laptop \"loan_month\" found."
                print(month_error)
                logging.error(month_error)
                abort = True
            if loan_date is None:
                date_error = "Error: no entry for laptop \"loan_date\" found."
                print(date_error)
                logging.error(date_error)
                abort = True
            if engineer is None:
                engineer_error = "Error: no entry for laptop \"engineer_id\" found."
                print(engineer_error)
                logging.error(engineer_error)
                abort = True
            if abort:
                abort_msg = f"Aborted adding laptop {json_object} to the database"
                print(abort_msg)
                logging.error(abort_msg)
                return
            new_laptop = self.laptop_utils.add_laptop_db(model, date(loan_year, loan_month, loan_date), engineer)
            success_msg = "New laptop info:\n" + str(new_laptop) + "\nSuccessfully added new laptop to the database!"
            print(success_msg)
            logging.info(success_msg)

        elif data_type == "contact_details":
            logging.info("Attempting to add new contact details from JSON object.")
            phone_number = json_object.get('phone_number', None)
            address = json_object.get('address', None)
            engineer = json_object.get('engineer', None)
            abort = False
            if phone_number is None:
                phone_error = "Error: no entry for contact details \"phone_number\" found."
                print(phone_error)
                logging.error(phone_error)
                abort = True
            if address is None:
                addr_error = "Error: no entry for contact details \"address\" found."
                print(addr_error)
                logging.error(addr_error)
                abort = True
            if engineer is None:
                engin_error = "Error: no entry for contact details \"engineer\" found."
                print(engin_error)
                logging.error(engin_error)
                abort = True
            engin = self.engin_utils.read_engineer_by_name(engineer)
            if engin is None:
                engin_error = f"Error: engineer {engineer} does not exist in the database. Contact details cannot be added without an existing engineer."
                print(engin_error)
                logging.error(engin_error)
                abort = True
            if abort:
                abort_msg = f"Aborted adding contact details {json_object} to the database."
                print(abort_msg)
                logging.error(abort_msg)
                return
            new_contact = self.contact_utils.add_contact_details_db(phone_number, address, engineer)
            if new_contact is None:
                dup_error = f"Contact details with phone number {phone_number} already exists." + \
                            f"\nAborted adding contact details {json_object} to the database."
                print(dup_error)
                logging.info(dup_error)
                return
            success_msg = "New contact details info:\n" + str(new_contact) + "\nSuccessfully added new contact details to the database!"
            print(success_msg)
            logging.info(success_msg)

        else:
            data_type_error = f"\"data_type\" entry value of {data_type} is not recognized.\n" + \
                              "\"data_type\" should be one of: \"vehicle\", \"engineer\", \"laptop\", \"contact_details\"\n" + \
                              f"Aborted adding entry {json_object} to the database."
            print(data_type_error)
            logging.error(data_type_error)
            return

    def display_interface(self):
        tables = ["Vehicles", "Engineers", "Laptops", "Contact Details"]
        function_map = [
            [self.add_vehicle, self.delete_vehicle, self.query_vehicles],
            [self.add_engineer, self.delete_engineer, self.query_engineers],
            [self.add_laptop, self.delete_laptop, self.query_laptops],
            [self.add_contact, self.delete_contact, self.query_contacts],
            self.insert_from_json
        ]
        print("\nWelcome to your Database Utilities!")
        print("1. Vehicles")
        print("2. Engineers")
        print("3. Laptops")
        print("4. Contact Details")
        print("5. Insert data from a JSON file")
        print("6: Exit Database Utilities")

        table_choice = get_digit_choice("Select a table, JSON insertion, or Exit:",
                                        "Invalid selection. Please enter a digit corresponding to the desired table, JSON insert, or Exit.",
                                        1, 7)
        table_choice -= 1

        if table_choice == 5:
            return False

        if table_choice == 4:
            self.insert_from_json()
            return True

        print(f"\nWhat would you like to do in the {tables[table_choice]} table?")
        print(f"1. Add a(n) {tables[table_choice][0:-1]}")
        print(f"2. Delete a(n) {tables[table_choice][0:-1]}")
        print(f"3. Read from the {tables[table_choice]} table")
        print("4. Choose a different table to work with")

        action_choice = get_digit_choice(f"Select an action to perform on the {tables[table_choice]} table:",
                                         "Invalid selection. Please enter a digit corresponding to the desired action.",
                                         1, 5)
        action_choice -= 1

        if action_choice == 3:
            # tell main() to restart the interface
            return True

        # call the chosen function from the function map
        function_map[table_choice][action_choice]()

        exit_choice = get_yes_no_choice("Exit the utilities script? (Y/N):")

        # returns false if user chose y to exit, tells main to stop running interface()
        return exit_choice.lower() != "y"